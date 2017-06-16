#! python 3

# HTML converter version 1.2 - take inDesign-generated .DOCX files  and
# convert it into plaintext HTML ready to upload to the BrooklynRail.org CMS.

# you can run this file by opening Terminal and typing
#   python3 FILE PATH
#   replacing FILE PATH with the path to this BR-HtmlConverter-1.2.py file
#   NOTE: you can drag-and-drop the file into the terminal window to get the file path

#   NOTE: if you get error messages when you first try to run the program, you likely do not have all of the necessary Python modules installed. You can find directions for installing these at the links below.
#   Don't be discouraged! Even if you are not familiar with the command line, you can install these modules with just a few simple lines of code. Note that when installing modules, you should type in python3 instead of python and pip3 instead of pip.

# version 1.2 changes
    # improved command prompts and code-commenting to make the script easier to use and maintain

# version 1.1 changes
    # changed inline styles to new CSS classes


import docx, os, mammoth, re, send2trash



print('''
Welcome to the Brooklyn Rail HTML Converter v1.2
This is a command-line script that will convert a Word document exported from InDesign into an HTML file.
This file will contain all of the articles in an issue, ready to be cut-and-pasted into the Brooklyn Rail CMS.
This converter is helpful, but it is NOT PERFECT. It may not complete all of the necessary formatting.
You must take care to review each article before going live online.\n
You can edit the code of this script to include new InDesign Paragraph and Character styles.\n
NOTES:
- !!! You MUST remove all hyperlinks from the Word document before using this program, or text will go missing. !!!
    You can do this by selecting all text in Word and hitting CMD + SHIFT + F9
- !!! You MUST do a find-and-replace for Italics and Interview Names in the InDesign file before exporting the text. See tutorial for instructions. !!!
- If this script isn't working correctly, you may need to install the Mammoth python module. Instructions here: https://pypi.python.org/pypi/mammoth
- You may also need to install the send2trash module: https://pypi.python.org/pypi/Send2Trash
- You may also need to install the docx module: https://python-docx.readthedocs.io/en/latest/user/install.html
\n''')

input('\nHit ENTER to begin')

print('Enter the path to the text file you wish to convert:')
#open the BR text file [or folder]
issueDocxPath = input()
issueDocxPath = issueDocxPath.replace("\\ ", " ")

print('\nEnter a name for the converted HTML file:')

htmlPath = input() + '.html'

input('\nHit ENTER to begin conversion')

print('Document opened')

issueDocument = docx.Document(issueDocxPath)

styles = issueDocument.styles

numParagraphs = len(issueDocument.paragraphs)

def convertRunItalics(paragraph): #converts italics within style replacements
    for run in paragraph.runs:
            if run.font.italic == True:
                run.text = '<em>' + run.text + '</em>'
            elif 'Caption Italic' == str(run.style.name):
                run.text = '<em>' + run.text + '</em>'
            elif 'Semibold italic' == str(run.style.name):
                run.text = '<em>' + run.text + '</em>'


def convertParaStyle(styleName, openTag, closeTag):
    if styleName == str(paragraph.style.name):
        convertRunItalics(paragraph) #convert italics to em tags
        paragraph.text = openTag + paragraph.text + closeTag

def convertCharStyle(charStylesList):
    for run in paragraph.runs:
        for i in range(len(charStylesList)):
            if charStylesList[i][0] == str(run.style.name):
                run.text = charStylesList[i][1] + run.text + charStylesList[i][2]

def deleteStyleText(deleteStylesList):
    for i in range(len(deleteStylesList)):
        if deleteStylesList[i] == str(paragraph.style.name):
            paragraph.text = ''
        for run in paragraph.runs:
            if deleteStylesList[i] == str(run.style.name):
                run.text = ''

def makeCaps(styleName):
    for run in paragraph.runs:
        if str(run.style.name) == styleName:
            run.text = run.text.upper()

#some paragraph and character styles of text to delete in the document
deleteStylesList = ['pull quote 1 base (Pull Quotes)',                      #delete pull quotes
                    'IGNORE ME: RUNNING FOOTER (IGNORE)',                   #delete section footers
                    'TOC / BODY (TOC paragraph)',                           #delete ToC
                    'REGULAR NO RULE (TOC paragraph)',                      #delete ToC
                    'SECTION_RULEABOVE (TOC paragraph)',                    #delete ToC
                    'masthead_credits (TOC paragraph)']                     #delete masthead

#character styles to tag
#(Style Name, Open Tag, Close Tag)
charStylesList = [
                    ('Endnote Reference1', '<sup>', '</sup>'),                #superscripts
                    ('Endnote Reference', '<sup>', '</sup>'),                 #superscripts
                    ('byline B', '#BYLINE#', ''),                             #bylines [code formatting]
                    ('byline-caps-light', '#BYLINE#', '')                     #bylines [code formatting]
                    ]


#loop over paragraphs finding style, replace text within that style to surround with tag
print('Deleting unwanted text and tagging styles')
for paragraph in issueDocument.paragraphs:

    deleteStyleText(deleteStylesList)
    convertCharStyle(charStylesList)
    makeCaps('BIOLINE / NAME / TRADE GOTHIC BOLD')

    convertParaStyle('SPAN COLUMN ARTSEEN / WRITER (ARTSEEN)', '#BYLINE#', '')                           #bylines [code formatting]
    convertParaStyle('10/12 Minion-Long Quote (10/12 Minion Body justified)', '#BLOCKQUOTE#', '')        #blockquote general
    convertParaStyle('Theater Dialogue in Interview (THEATER)', '#BLOCKQUOTE#', '')                      #blockquote THEATER
    convertParaStyle('Theater Dialogue (THEATER)', '#BLOCKQUOTE#', '')                                   #blockquote THEATER
    convertParaStyle('10/12 Minion-Long Quote no break (10/12 Minion Body justified)', '#BQPOETRY#', '') #blockquote poetry
    convertParaStyle('Theater Dialogue CHARACTER (THEATER)', '#THEATERNAME#', '')                        #quote names THEATER
    convertParaStyle('april_letters_FROM (TOC paragraph)', '#LGQUOTENAME#', '#ENDLGQUOTENAME#')                   #quote attribution large
    convertParaStyle('10/12 Minion / Quote / attrib (10/12 Minion Body justified)', '#QUOTENAME#', '')            #quote attribution
    convertParaStyle('10/12 Minion-Long Quote attribution (10/12 Minion Body justified)', '#QUOTENAME#', '')      #quote attribution
    convertParaStyle('10/12 Minion / Quote / attrib longquote (10/12 Minion Body justified)', '#QUOTENAME#', '')  #quote attribution
    convertParaStyle('Paragraph Style 1 (BOOKS)', '#BKHEAD#', '#ENDBKHEAD#')                  #BOOKS headers
    convertParaStyle('COLUMN BOOK INFO ONE LINE (BOOKS)', '#BKHEAD#', '#ENDBKHEAD#')                  #BOOKS headers
    convertParaStyle('Title B (10/12 Minion Body justified)', '#TITLE#', '#ENDTITLE#')                                          #titles, [code formatting] general
    convertParaStyle('ARTBOOKS NEW TITLE regular (ARTBOOKS)', '#TITLE#', '#ENDTITLE#')                                          #titles, [code formatting] ARTBOOKS **NEW
    convertParaStyle('ARTBOOKS NEW TITLE ITALIC (ARTBOOKS)', '#TITLE#', '#ENDTITLE#')                                           #titles, [code formatting] ARTBOOKS **NEW
    convertParaStyle('title-lowercase-light (DANCE)', '#TITLE#', '#ENDTITLE#')                                                  #titles, [code formatting] DANCE
    convertParaStyle('36pt- (DANCE)', '#TITLE#', '#ENDTITLE#')                                                                  #titles, [code formatting] DANCE
    convertParaStyle('VERBATIM TITLE (VERBATIM)', '#TITLE#', '#ENDTITLE#')                                                      #titles, [code formatting] VERBATIM
    convertParaStyle('minion Title interviewee (ARTSEEN)', '#TITLE#', '#ENDTITLE#')                                             #titles, [code formatting] ARTSEEN
    convertParaStyle('SPAN COLUMN ARTSEEN / TITLE ROMAN (ARTSEEN)', '#TITLEAS#', '#ENDTITLEAS#')                                #titles, [code formatting] ARTSEEN
    convertParaStyle('SPAN COLUMN ARTSEEN / TITLE ITALIC (ARTSEEN)', '#TITLEASITAL#', '#ENDTITLEASITAL#')                       #titles, [code formatting] ARTSEEN italic
    convertParaStyle('10/12 Minion-Drop Cap (10/12 Minion Body justified)', '#DROPCAP#', '#ENDDROPCAP#')                     #drop cap fix 1char
    convertParaStyle('10/12 Minion-Drop Cap quotations (2 char) (10/12 Minion Body justified)', '#DROPCAP#', '#ENDDROPCAP#') #drop cap fix 2char
    convertParaStyle('BIOLINE 7/9 TradeGothic (10/12 Minion Body justified)', '#BIO#', '#ENDBIO#')      #bios, [code formatting]
    convertParaStyle('Image Captions', '#IMGCAP#', '#ENDIMGCAP#')                                       #image captions
    convertParaStyle('PARA BREAK - ASTERIK (10/12 Minion Body justified)', '<p align="center">', '')    #center break characters
    convertParaStyle('Endnotes with rule 2 USE THIS  (10/12 Minion Body justified)', '#ENDHEAD#', '#ENDENDHEAD#')                         #endnotes header
    convertParaStyle('10/12 Minion-Announcements & Endnotes list 9/11 continued USE THIS (10/12 Minion Body justified)', '<li>', '</li>') #endnotes list
    convertParaStyle('10/12 Minion-Announcements & Endnotes (10/12 Minion Body justified)', '#ANNOUNCETOP#', '')     #announcements top
    convertParaStyle('10/12 Minion-Announcements & Endnotes within (10/12 Minion Body justified)', '#ANNOUNCE#', '') #announcements within
    convertParaStyle('Theater Endnotes (10/12 Minion Body justified)', '#ANNOUNCETOP#', '')                          #announcements top THEATER
    convertParaStyle('MUSIC VENUE copy (MUSIC)', '#VENUE#', '')     #venue MUSIC
    convertParaStyle('2DANCE VENUE (DANCE)', '#VENUE#', '')         #venue DANCE
    convertParaStyle('venue-new (DANCE)', '#VENUE#', '')            #venue DANCE
    convertParaStyle('ARTSEEN_LEFT_VENUE (ARTSEEN)', '#VENUE#', '') #venue ARTSEEN
    convertParaStyle('Subhead Minion (10/12 Minion Body justified:Subheads)', '#SUBTITLE14#', '')   #subtitle general, 14pt
    convertParaStyle('POEM TITLE (SPANNING FICTION & POETRY)', '#SUBTITLE14CENT#', '')              #subtitle FICTION, 14pt centered
    convertParaStyle('SUBHEAD CHAPTERS nospan (SPANNING FICTION & POETRY)', '#SUBTITLE14CENT#', '') #subtitle FICTION, 14pt centered
    convertParaStyle('ARTSEEN_LEFT_SHOWTITLE (ARTSEEN)', '#subtitleBoldCent#', '') #subtitle ARTSEEN shows, bold italic centered
    convertParaStyle('ARTSEEN_LEFT_TITLE ITALSHOW (ARTSEEN)', '#subtitleBoldCent#', '') #subtitle ARTSEEN shows, bold italic centered

issueDocument.save('converted.docx')

#convert to html
print('Passing to mammoth')
with open('converted.docx', 'rb') as docx_file:
    result = mammoth.convert_to_html(docx_file)
    html = result.value
    messages = result.messages
print('File converted to html')

send2trash.send2trash('converted.docx')

#------------------------------------------------------------------------------------------------------
#start find and replace
#------------------------------------------------------------------------------------------------------

print('Replacing Special Characters')

html = html.replace('</p><p>', '</p>\n\n<p>') #add new lines between paragraphs to aid readability


#------------------------------------------------------------------------------------------------------
#convert <, >, and " from entities to characters to have functional html tags
html = html.replace('&lt;', '<')
html = html.replace('&gt;', '>')
html = html.replace('&quot;', '"')

#convert unicode special characters to HTML entities
html = html.replace('\u201c', '&ldquo;') #left double quote
html = html.replace('\u201d', '&rdquo;') #right double quote
html = html.replace('\u2018', '&lsquo;') #left single quote
html = html.replace('\u2019', '&rsquo;') #right single quote
html = html.replace('\u2026', '&#8230;') #ellipsis character
html = html.replace('\u2013', '&#8211;') #n-dash
html = html.replace('\u2014', '&#8212;') # em-dash
html = html.replace('\u00a0', '&#160;') #no-break space
html = html.replace('\u00a9', '&#169;') #copyright symbol (c)
html = html.replace('\u00c0', '&#192;') #grave A
html = html.replace('\u00c1', '&#193;') #acute A
html = html.replace('\u00c2', '&#194;') #circumflex A
html = html.replace('\u00c3', '&#195;') #tilde A
html = html.replace('\u00c4', '&#196;') #umlaut A
html = html.replace('\u00c5', '&#197;') #bolle A
html = html.replace('\u00c6', '&#198;') #AE
html = html.replace('\u00c7', '&#199;') #cedilla C
html = html.replace('\u00c8', '&#200;') #grave E
html = html.replace('\u00c9', '&#201;') #acute E
html = html.replace('\u00ca', '&#202;') #circumflex E
html = html.replace('\u00cb', '&#203;') #umlaut E
html = html.replace('\u00cc', '&#204;') #grave I
html = html.replace('\u00cd', '&#205;') #acute I
html = html.replace('\u00ce', '&#206;') #circumflex I
html = html.replace('\u00cf', '&#207;') #umlaut I
html = html.replace('\u00d0', '&#208;') #D with line
html = html.replace('\u00d1', '&#209;') #tilde N
html = html.replace('\u00d2', '&#210;') #grave O
html = html.replace('\u00d3', '&#211;') #acute O
html = html.replace('\u00d4', '&#212;') #circumflex O
html = html.replace('\u00d5', '&#213;') #tilde O
html = html.replace('\u00d6', '&#214;') #umlaut O
html = html.replace('\u00d7', '&#215;') #x symbol
html = html.replace('\u00d8', '&#216;') #O with line
html = html.replace('\u00d9', '&#217;') #grave U
html = html.replace('\u00da', '&#218;') #acute U
html = html.replace('\u00db', '&#219;') #circumflex U
html = html.replace('\u00dc', '&#220;') #umlaut U
html = html.replace('\u00dd', '&#221;') #Y acute
html = html.replace('\u00de', '&#222;') #P thingy
html = html.replace('\u00df', '&#223;') #B thingy
html = html.replace('\u00e0', '&#224;') #grave a
html = html.replace('\u00e1', '&#225;') #acute a
html = html.replace('\u00e2', '&#226;') #circumflex a
html = html.replace('\u00e3', '&#227;') #tilde a
html = html.replace('\u00e4', '&#228;') #umlaut a
html = html.replace('\u00e5', '&#229;') #bolle a
html = html.replace('\u00e6', '&#230;') #ae
html = html.replace('\u00e7', '&#231;') #cedilla c
html = html.replace('\u00e8', '&#232;') #grave e
html = html.replace('\u00e9', '&#233;') #acute e
html = html.replace('\u00ea', '&#234;') #circumflex e
html = html.replace('\u00eb', '&#235;') #umlaut e
html = html.replace('\u00ec', '&#236;') #grave i
html = html.replace('\u00ed', '&#237;') #acute i
html = html.replace('\u00ee', '&#238;') #circumflex i
html = html.replace('\u00ef', '&#239;') #umlaut i
html = html.replace('\u00f0', '&#240;') #d thingy
html = html.replace('\u00f1', '&#241;') #tilde n
html = html.replace('\u00f2', '&#242;') #grave o
html = html.replace('\u00f3', '&#243;') #acute o
html = html.replace('\u00f4', '&#244;') #circumflex o
html = html.replace('\u00f5', '&#245;') #tilde o
html = html.replace('\u00f6', '&#246;') #umlaut o
html = html.replace('\u00f7', '&#247;') #division sign
html = html.replace('\u00f8', '&#248;') #o with line
html = html.replace('\u00f9', '&#249;') #grave u
html = html.replace('\u00fa', '&#250;') #acute u
html = html.replace('\u00fb', '&#251;') #circumflex u
html = html.replace('\u00fc', '&#252;') #umlaut u
html = html.replace('\u00fd', '&#253;') #acute y
html = html.replace('\u00fe', '&#254;') #p thingy
html = html.replace('\u00ff', '&#255;') #umlaut y
html = html.replace('\u014d', '&#333;') #macron o
html = html.replace('\u2022', '&#8226;') #bullet

#invisible spaces
html = html.replace('\u2063', '') # &#8291;
html = html.replace('\u1680', '') #ogham space mark
html = html.replace('\u180e', '') #mongolian vowel separator
html = html.replace('\u2000', '') #en quad
html = html.replace('\u2001', '') #em quad
html = html.replace('\u2002', '') #en space
html = html.replace('\u2003', '') #em space
html = html.replace('\u2004', '') #three-per-em space
html = html.replace('\u2005', '') #four-per-em space
html = html.replace('\u2006', '') #six-per-em space
html = html.replace('\u2007', '') #figure space
html = html.replace('\u2008', '') #punctuation space
html = html.replace('\u2009', '') #thin space
html = html.replace('\u200A', '') #hair space
html = html.replace('\u200B', '') #zero width space
html = html.replace('\u202F', '') #narrow no-break space
html = html.replace('\u205F', '') #medium mathematcal space
html = html.replace('\u3000', '') #ideographic space
html = html.replace('\uFEFF', '') #zero width no-break space

#fractions
html = html.replace('\uF00bc', '1/4') #vulgar fraction 1/4
html = html.replace('\uF00bd', '1/2') #vulgar fraction 1/2
html = html.replace('\uF00be', '3/4') #vulgar fraction
#------------------------------------------------------------------------------------------------------

print('Inserting HTML tags')
#remove duplicate tags
html = html.replace('<p><p align="center">', '<p align="center">')
html = html.replace('<p><p style="font-size:10pt;margin-left:20px;margin-top:-1em;">', '<p style="font-size:10pt;margin-left:20px;margin-top:-1em;">')


html = html.replace('<p>#BLOCKQUOTE#', '<p class="blockquote">')                    			#blockquote
html = html.replace('<p>#BQPOETRY#', '<p class="blockquote-poetry">')               			#blockquote poetry
html = html.replace('<p>#QUOTENAME#', '<p class="quote-attribution">')              			#quote attribution
html = html.replace('<p>#TITLE#', '\n\n\n\n\n<p>&nbsp;</p>\n<p>&nbsp;</p>\n\n\n\n\nARTICLE TITLE:\n')           #titles [code formatting]
html = html.replace('#ENDTITLE#</p>', '\n\n<br/><br/>\n\nARTICLE TEXT:')                                        #titles [code formatting]
html = html.replace('<p>#TITLEAS#', '\n\n\n\n\n<p>&nbsp;</p>\n<p>&nbsp;</p>\n\n\n\n\nARTICLE TITLE:\n')         #titles ARTSEEN [code formatting]
html = html.replace('#ENDTITLEAS#</p>', '\n\n<br/><br/>\n\nARTICLE TEXT:')                                      #titles ARTSEEN [code formatting]
html = html.replace('<p>#TITLEASITAL#', '\n\n\n\n\n<p>&nbsp;</p>\n<p>&nbsp;</p>\n\n\n\n\nARTICLE TITLE:\n<em>') #titles ARTSEEN ital [code formatting]
html = html.replace('#ENDTITLEASITAL#</p>', '</em>\n\n<br/><br/>\n\nARTICLE TEXT:')                             #titles ARTSEEN ital [code formatting]
html = html.replace('<p>#VENUE#', '<p class="artseen-info">')                       			#venue header
html = html.replace('<p>#BIO#', '\n\n--------------------\n<br/>\nBIO:<br/>\n\n')   			#bios [code formatting]
html = html.replace('#ENDBIO#</p>', '\n\n')                                        				#bios [code formatting]
html = html.replace('<p>#IMGCAP#', '\n\nIMAGE CAPTION:\n')      								#image captions [code formatting]
html = html.replace('#ENDIMGCAP#</p>', '\n\n')                  								#image captions [code formatting]
html = html.replace('<p>#ENDHEAD#', '<hr />\n\n<p class="endnote">')        					#endnotes header
html = html.replace('#ENDENDHEAD#</p>', '</p>\n<ol class="endnote">\n')     					#endnotes header
html = html.replace('<p><li>', '<li>')                          								#endnotes
html = html.replace('</li></p>', '</li></ol>')                  								#endnotes
html = html.replace('</ol>\n\n<li>', '\n<li>')                  								#endnotes
html = html.replace('#ENDDROPCAP#</p>\n\n<p>#DROPCAP#', '')     								#drop cap FIRST (middle duplicate removal)
html = html.replace('<p>#DROPCAP#', '<p>')                      								#drop cap SECOND (top)
html = html.replace('#ENDDROPCAP#', '</p>\n\n<!--insert first image here-->\n<p>')              #drop cap THIRD (end)
html = html.replace('<p>#THEATERNAME#', '<p class="theater-name">')    							#THEATER quote names
html = html.replace('<p>#ANNOUNCETOP#', '<hr />\n<p class="endnote">') 							#announcements top
html = html.replace('<p>#ANNOUNCE#', '<p class="endnote">')            							#announcements within
html = html.replace('<p>#BKHEAD#', '<p class="book-info">\n')                              		#BOOKS headers
html = html.replace('#ENDBKHEAD#', '')                                                 			#BOOKS headers
html = html.replace('<p>#SUBTITLE14#', '<p class="subtitle">')                          		#subtitle 14 pt
html = html.replace('<p>#SUBTITLE14CENT#', '<p class="subtitle-cent">')       					#subtitle 14 pt centered
html = html.replace('<p>#subtitleBoldCent#', '<p class="subtitle-cent" style="font-style: italic;">')  #subtitle bold italic centered
html = html.replace('<p>#LGQUOTENAME#', '<p style="font-size: 14pt;" align="right"><em>') 		#large quote attributions, 14 pt
html = html.replace('#ENDLGQUOTENAME#', '</em>') 												#large quote attributions, 14 pt


#delete any leftover tags and other special cases, such as layout cleanup

html = html.replace('--------------------\n<br/>\nBIO:<br/>\n\n\n', '')
html = html.replace('ARTICLE TITLE:\nin conversation', 'ARTICLE TYPE: in conversation\n<br /><br />\nARTICLE TITLE: ') #cleans up "in conversation" article type header
html = html.replace('BYLINE:\n \n<br/>\nBYLINE:\n<br />','BYLINE: ')							#duplicate "BYLINE" header
html = html.replace('<li></li>', '') 															#delete empty list items
html = html.replace('<p></p>', '') 																#delete empty paragraphs
html = html.replace('ARTICLE TEXT:\n\n\n', '') 													#delete extra "article text" headers in the html [code formatting]
html = html.replace('#BYLINE#', '\n<br/><br/>\nBYLINE:\n') 										#bylines [code formatting]



#------------------------------------------------------------------------------------------------------
#convert URLs to hyperlinks

print('Hyperlinking URLs')

urlRegex1 = re.compile(r'((http|ftp|https)(://)([\w_-]+(?:(?:\.[\w_-]+)+))([\w.,@?^=%&:/~+#-]*[\w@?^=%&/~+#-])?)')
urlRegex2 = re.compile(r'( |\()((\w|\.|-)*(\.com|\.org|\.net)(\w|\b\.\b|/|-)*)') #need GROUP 2 out of this regex, as it checks for and includes white space at the beginning

emailRegex = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,4}')

urls1 = urlRegex1.findall(html)
urls2 = urlRegex2.findall(html)
emails = emailRegex.findall(html)

for i in range(len(urls1)): 	#hyperlink the urls that have http
    url = urls1[i][0]
    html = html.replace(url, '<a href="' + url + '">' + url + '</a>')

for i in range(len(urls2)): 	#hyperlink the urls missing http
    url = urls2[i][1]
    html = html.replace(url, '<a href="http://' + url + '">' + url + '</a>')

for i in range(len(emails)): 	#hyperlink the emails
    email = emails[i]
    html = html.replace(email, '<a href="mailto:' + email + '">' + email + '</a>')


#------------------------------------------------------------------------------------------------------
#end find and replace
#------------------------------------------------------------------------------------------------------



#save the file

htmlFile = open(htmlPath, 'w', encoding='utf8')
htmlFile.write(html)
htmlFile.close()

print('Done. ' + htmlPath + ' output to Home directory' )
